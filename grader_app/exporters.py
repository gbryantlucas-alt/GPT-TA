from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from grader_app.models import BatchSession


def export_canvas_csv(session: BatchSession, out_dir: str) -> str:
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    rows = []
    for sid, essay in session.essays.items():
        row = {"student_id": sid, "overall_grade": essay.overall_grade}
        for cat in essay.category_scores:
            row[f"{cat.dimension}_score"] = cat.score
            row[f"{cat.dimension}_label"] = cat.label
        rows.append(row)
    df = pd.DataFrame(rows)
    out_path = Path(out_dir) / f"{session.session_id}_canvas_export.csv"
    df.to_csv(out_path, index=False)
    return str(out_path)


def export_student_feedback_files(session: BatchSession, out_dir: str) -> list[str]:
    out = Path(out_dir) / "feedback"
    out.mkdir(parents=True, exist_ok=True)
    files: list[str] = []
    for sid, essay in session.essays.items():
        doc = Document()
        doc.add_heading(f"Feedback Report - {sid}", level=1)
        doc.add_heading("Essay Summary", level=2)
        doc.add_paragraph(essay.summary)

        doc.add_heading("Assignment Compliance", level=2)
        for item in essay.compliance:
            doc.add_paragraph(f"[{item.status.upper()}] {item.requirement}: {item.note}")

        doc.add_heading("Rubric Scores & Feedback", level=2)
        for score in essay.category_scores:
            doc.add_paragraph(f"{score.dimension}: {score.score} ({score.label})")
            doc.add_paragraph(score.feedback)

        doc.add_heading("Human Judgment Flags", level=2)
        for ann in essay.annotations:
            doc.add_paragraph(f"{ann.dimension} | Excerpt: {ann.excerpt}\nTeacher review question: {ann.question}")

        doc.add_heading("Overall Grade", level=2)
        doc.add_paragraph(f"{essay.overall_grade}")
        doc.add_paragraph(essay.overall_note)
        if essay.ai_suspicion_note:
            doc.add_heading("AI Usage Signals", level=2)
            doc.add_paragraph(f"Score: {essay.ai_suspicion_score}")
            doc.add_paragraph(essay.ai_suspicion_note)

        file_path = out / f"{sid}_feedback.docx"
        doc.save(file_path)
        files.append(str(file_path))
    return files
